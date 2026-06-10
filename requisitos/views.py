import json
import io
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.utils import timezone
from .models import Requisito, MensagemChat, TemplateElicitacao
from .forms import RequisitoForm, MensagemChatForm

# ── CRUD ──────────────────────────────────────────────────────────────────────

@login_required
def lista(request):
    requisitos = Requisito.objects.all()
    tipo = request.GET.get('tipo', '')
    status = request.GET.get('status', '')
    if tipo:
        requisitos = requisitos.filter(tipo=tipo)
    if status:
        requisitos = requisitos.filter(status=status)
    return render(request, 'requisitos/lista.html', {
        'requisitos': requisitos,
        'tipo_filtro': tipo,
        'status_filtro': status,
    })

@login_required
def criar(request):
    template_id = request.GET.get('template')
    template = None
    if template_id:
        template = TemplateElicitacao.objects.filter(pk=template_id).first()

    if request.method == 'POST':
        form = RequisitoForm(request.POST)
        if form.is_valid():
            req = form.save(commit=False)
            req.criado_por = request.user
            req.save()
            form.save_m2m()
            messages.success(request, f'Requisito {req.codigo} criado com sucesso!')
            return redirect('requisitos:lista')
    else:
        form = RequisitoForm()
    return render(request, 'requisitos/form.html', {
        'form': form,
        'titulo': 'Novo Requisito',
        'template': template,
        'templates': TemplateElicitacao.objects.prefetch_related('perguntas').all(),
    })

@login_required
def editar(request, pk):
    requisito = get_object_or_404(Requisito, pk=pk)
    if request.method == 'POST':
        form = RequisitoForm(request.POST, instance=requisito)
        if form.is_valid():
            form.save()
            messages.success(request, 'Requisito atualizado com sucesso!')
            return redirect('requisitos:lista')
    else:
        form = RequisitoForm(instance=requisito)
    return render(request, 'requisitos/form.html', {
        'form': form,
        'titulo': 'Editar Requisito',
        'obj': requisito,
    })

@login_required
def excluir(request, pk):
    requisito = get_object_or_404(Requisito, pk=pk)
    if request.method == 'POST':
        requisito.delete()
        messages.success(request, 'Requisito removido.')
        return redirect('requisitos:lista')
    return render(request, 'requisitos/confirmar_exclusao.html', {'obj': requisito})

# ── HISTÓRICO ─────────────────────────────────────────────────────────────────

@login_required
def historico(request, pk):
    requisito = get_object_or_404(Requisito, pk=pk)
    versoes = requisito.history.all()
    return render(request, 'requisitos/historico.html', {
        'requisito': requisito,
        'versoes': versoes,
    })

# ── CHAT ──────────────────────────────────────────────────────────────────────

@login_required
def chat(request, pk):
    requisito = get_object_or_404(Requisito, pk=pk)
    mensagens = requisito.mensagens.select_related('usuario').prefetch_related('likes').all()

    if request.method == 'POST':
        form = MensagemChatForm(request.POST)
        if form.is_valid():
            msg = form.save(commit=False)
            msg.requisito = requisito
            msg.usuario = request.user
            msg.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'id': msg.id,
                    'usuario': msg.usuario.username,
                    'mensagem': msg.mensagem,
                    'data': msg.data.strftime('%d/%m/%Y %H:%M'),
                    'total_likes': 0,
                    'liked': False,
                })
            return redirect('requisitos:chat', pk=pk)
    else:
        form = MensagemChatForm()

    msgs_data = []
    for m in mensagens:
        msgs_data.append({
            'id': m.id,
            'usuario': m.usuario.username,
            'mensagem': m.mensagem,
            'data': m.data.strftime('%d/%m/%Y %H:%M'),
            'total_likes': m.total_likes,
            'liked': request.user in m.likes.all(),
            'proprio': m.usuario == request.user,
        })

    return render(request, 'requisitos/chat.html', {
        'requisito': requisito,
        'mensagens': mensagens,
        'msgs_json': json.dumps(msgs_data),
        'form': form,
    })

@login_required
@require_POST
def like_mensagem(request, pk):
    mensagem = get_object_or_404(MensagemChat, pk=pk)
    if request.user in mensagem.likes.all():
        mensagem.likes.remove(request.user)
        liked = False
    else:
        mensagem.likes.add(request.user)
        liked = True
    return JsonResponse({'total_likes': mensagem.total_likes, 'liked': liked})

@login_required
def buscar_mensagens(request, pk):
    """Polling para colaboracao em tempo real (atualiza a cada 10s)"""
    requisito = get_object_or_404(Requisito, pk=pk)
    since_id = request.GET.get('since', 0)
    mensagens = requisito.mensagens.filter(id__gt=since_id).select_related('usuario').prefetch_related('likes')
    data = [{
        'id': m.id,
        'usuario': m.usuario.username,
        'mensagem': m.mensagem,
        'data': m.data.strftime('%d/%m/%Y %H:%M'),
        'total_likes': m.total_likes,
        'liked': request.user in m.likes.all(),
        'proprio': m.usuario == request.user,
    } for m in mensagens]
    return JsonResponse({'mensagens': data})

# ── MATRIZ DE RASTREABILIDADE ─────────────────────────────────────────────────

@login_required
def matriz_rastreabilidade(request):
    requisitos = Requisito.objects.prefetch_related('stakeholders').all()
    from stakeholders.models import Stakeholder
    stakeholders = Stakeholder.objects.all()
    return render(request, 'requisitos/matriz.html', {
        'requisitos': requisitos,
        'stakeholders': stakeholders,
    })

# ── EXPORTAÇÕES ───────────────────────────────────────────────────────────────

@login_required
def exportar_json(request):
    requisitos = Requisito.objects.prefetch_related('stakeholders').all()
    data = [{
        'codigo': r.codigo,
        'titulo': r.titulo,
        'descricao': r.descricao,
        'tipo': r.get_tipo_display(),
        'prioridade': r.get_prioridade_display(),
        'status': r.get_status_display(),
        'stakeholders': [s.nome for s in r.stakeholders.all()],
        'criado_por': r.criado_por.username if r.criado_por else None,
        'criado_em': r.criado_em.strftime('%d/%m/%Y %H:%M'),
    } for r in requisitos]
    response = JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False, 'indent': 2})
    response['Content-Disposition'] = 'attachment; filename="requisitos.json"'
    return response

@login_required
def exportar_markdown(request):
    requisitos = Requisito.objects.prefetch_related('stakeholders').all()
    lines = ['# TraceMind — Requisitos\n']
    for r in requisitos:
        lines.append(f'## {r.codigo} — {r.titulo}')
        lines.append(f'- **Tipo:** {r.get_tipo_display()}')
        lines.append(f'- **Prioridade:** {r.get_prioridade_display()}')
        lines.append(f'- **Status:** {r.get_status_display()}')
        lines.append(f'- **Stakeholders:** {", ".join(s.nome for s in r.stakeholders.all()) or "Nenhum"}')
        lines.append(f'\n{r.descricao}\n')
    content = '\n'.join(lines)
    response = HttpResponse(content, content_type='text/markdown')
    response['Content-Disposition'] = 'attachment; filename="requisitos.md"'
    return response

@login_required
def exportar_pdf(request):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    # Título
    title_style = ParagraphStyle('title', parent=styles['Title'], fontSize=18, textColor=colors.HexColor('#1e3a5f'))
    story.append(Paragraph('TraceMind — Especificacao de Requisitos', title_style))
    story.append(Paragraph(f'Gerado em: {timezone.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal']))
    story.append(Spacer(1, 0.5*cm))

    requisitos = Requisito.objects.prefetch_related('stakeholders').all()

    for r in requisitos:
        # Cabeçalho do requisito
        story.append(Paragraph(f'{r.codigo} — {r.titulo}', styles['Heading2']))
        data = [
            ['Tipo', r.get_tipo_display(), 'Prioridade', r.get_prioridade_display()],
            ['Status', r.get_status_display(), 'Stakeholders', ', '.join(s.nome for s in r.stakeholders.all()) or 'Nenhum'],
        ]
        t = Table(data, colWidths=[3*cm, 5*cm, 3*cm, 6*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#e8edf5')),
            ('BACKGROUND', (2,0), (2,-1), colors.HexColor('#e8edf5')),
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('PADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(t)
        story.append(Paragraph(r.descricao, styles['Normal']))
        story.append(Spacer(1, 0.4*cm))

    doc.build(story)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="requisitos.pdf"'
    return response

@login_required
def exportar_docx(request):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Título
    titulo = doc.add_heading('TraceMind — Especificacao de Requisitos', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Gerado em: {timezone.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph()

    requisitos = Requisito.objects.prefetch_related('stakeholders').all()

    for r in requisitos:
        doc.add_heading(f'{r.codigo} — {r.titulo}', level=2)
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        headers = ['Tipo', r.get_tipo_display(), 'Prioridade', r.get_prioridade_display()]
        row2 = ['Status', r.get_status_display(), 'Stakeholders', ', '.join(s.nome for s in r.stakeholders.all()) or 'Nenhum']
        for i, val in enumerate(headers):
            table.rows[0].cells[i].text = val
        for i, val in enumerate(row2):
            table.rows[1].cells[i].text = val
        doc.add_paragraph(r.descricao)
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="requisitos.docx"'
    return response
